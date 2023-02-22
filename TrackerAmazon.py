# sourcery skip: merge-nested-ifs, use-fstring-for-concatenation
from bs4 import BeautifulSoup
import requests
import docx
from docx import Document
import os
import urllib.request
import shutil
import locale
import csv
import winsound
import sys
import tempfile
import shutil
import win32com.client
locale.setlocale(locale.LC_ALL, 'it_IT.UTF-8')


def create_shortcut(path, target, description=None):
    if sys.platform == 'win32':
        shell = win32com.client.Dispatch("WScript.Shell")
        shortcut = shell.CreateShortCut(path)
        shortcut.Targetpath = target
        if description:
            shortcut.Description = description
        shortcut.save()
    else:
        print("Creating shortcuts is only supported on Windows.")

if __name__ == '__main__':
    current_dir = os.path.dirname(os.path.abspath(__file__))
    exe_file = os.path.join(current_dir, "TrackerAmazon.exe")
    path = os.path.expanduser("~/Desktop/TrackerAmazon.lnk")
    target = exe_file
    description = "Tracker prodotti Amazon"
    create_shortcut(path, target, description)


"""
Questa funzione crea gli hyperlink su word, fonti:https://github.com/python-openxml/python-docx/issues/74#issuecomment-261169410

"""


def add_hyperlink(paragraph, url, text, color, underline):
    # sourcery skip: de-morgan, extract-duplicate-method
    """
    Questa è una funzione che inserisce un ipertesto all'interno di un oggetto paragrafo.

    :param paragraph: il paragrafo al quale stiamo aggiungendo l'ipertesto
    :param url: una stringa che contiene l'URL richiesto
    :param text: il testo visualizzato per l'URL
    :return: l'oggetto ipertesto.

    """

    # Questo ottiene l'accesso al file document.xml.rels e ottiene un nuovo valore di id di relazione
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Crea il tag w:hyperlink e aggiunge i valori necessari
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Crea un elemento w:r
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Crea un nuovo elemento w:rPr
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Aggiunge il colore se viene dato
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Rimuove la sottolineatura se viene richiesto
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Unisce tutti gli elementi xml insieme e aggiunge il testo richiesto all'elemento w:r
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


while(True):
    scelta_def = int(input("Premi 1 se vuoi fare una ricerca su Amazon con filtro prezzi, premi 2 se invece vuoi tracciare i prezzi dei tuoi prodotti, se vuoi uscire dal programma premi 0: "))



    if scelta_def == 1:
        #creo il documento word
        # sourcery skip: for-index-underscore, merge-nested-ifs, remove-redundant-continue, remove-str-from-fstring, sum-comprehension, switch, use-fstring-for-concatenation
        document = Document()
        
        #controlla che non esista tale cartella, nel caso la rimuove e la ricrea
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        temp_folder = os.path.join(desktop_path, "Immagini Temporanee")
        
        if "Immagini Temporanee" in os.listdir(desktop_path):
            #elimina i file nella cartella per poi eliminare la cartella stessa
            for filename in os.listdir(temp_folder):
                file_path = os.path.join(temp_folder, filename)
                try:
                    if os.path.isfile(file_path) or os.path.islink(file_path):
                        os.unlink(file_path)
                    elif os.path.isdir(file_path):
                        shutil.rmtree(file_path)
                except Exception as e:
                    print(f'Failed to delete {file_path}. Reason: {e}')
            os.rmdir(temp_folder)
        os.mkdir(temp_folder)



        

        #############################################################################################################

        ricerca= input("Cosa vuoi cercare su Amazon: ")
        filtro_prezzo = float(input("Inserisci il massimo valore in euro che vuoi spendere: "))
        HEADERS = ({'User-Agent':
                    'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:106.0) Gecko/20100101 Firefox/106.0'})
        URL = f"https://www.amazon.it/s?k={ricerca}"

        webpage = requests.get(URL, headers=HEADERS)

        soup = BeautifulSoup(webpage.text, "html.parser")

        block_div= soup.find_all("div", attrs={"class":"s-card-container s-overflow-hidden aok-relative puis-expand-height puis-include-content-margin puis s-latency-cf-section s-card-border"})


        #il seguente for cicla tutti i blocchi div trovati sopra per cercare i nomi dei prodotti, i prezzi, i link al prodotto e l'immagine da scaricare


        for block in range(len(block_div)):
                
                #scarico le immagini
                desktop_path = os.path.expanduser("~/Desktop/")
                folder = os.path.join(desktop_path, "Immagini Temporanee/")
                file_name = str(block_div[block].find("span", attrs={"class":'a-size-base-plus a-color-base a-text-normal'})).replace("<span class=\"a-size-base-plus a-color-base a-text-normal\">", "").replace("</span>", "").replace("\\", "").replace("/", "").replace(":", "").replace("*", "").replace("?", "").replace("<", "").replace(">", "").replace("|", "").replace("\"", "")+".jpg"
                file_path = os.path.join(folder, file_name)
                with open(file_path, 'wb') as f:
                    f.write(urllib.request.urlopen(str(block_div[block].find("img")["src"])).read())

                    
                """
                Le istruzioni che seguono inseriscono nel file word, il prezzo, il nome del prodotto, il link e l'immagine

                """
                if block_div[block].find("span", attrs={'class':'a-price-whole'}) is not None:
                    if locale.atof(str(block_div[block].find("span", attrs={'class':'a-price-whole'})).replace("<span class=\"a-price-whole\">", "").replace("</span>", ""))<=filtro_prezzo:
                            
                        document.add_heading(str(block_div[block].find("span", attrs={"class":'a-size-base-plus a-color-base a-text-normal'})).replace("<span class=\"a-size-base-plus a-color-base a-text-normal\">", "").replace("</span>", "").replace("\\", "").replace("/", "").replace(":", "").replace("*", "").replace("?", "").replace("<", "").replace(">", "").replace("|", "").replace("\"", ""), 2)          
                        document.add_paragraph(str(block_div[block].find("span", attrs={'class':'a-price-whole'})).replace("<span class=\"a-price-whole\">", "").replace("</span>", "")+" "+"€")           
                                
                        p = document.add_paragraph("")
                        hyperlink= add_hyperlink(p, "https://www.amazon.it"+str(block_div[block].find("a")['href']), "Link al prodotto", None, True)                
                        file_path = os.path.join(os.path.expanduser("~/Desktop"), "Immagini Temporanee", str(block_div[block].find("span", attrs={"class":'a-size-base-plus a-color-base a-text-normal'})).replace("<span class=\"a-size-base-plus a-color-base a-text-normal\">", "").replace("</span>", "").replace("\\", "").replace("/", "").replace(":", "").replace("*", "").replace("?", "").replace("<", "").replace(">", "").replace("|", "").replace("\"", "") + ".jpg")
                        document.add_picture(file_path)

                        # questa istruzione passa alla pagina successiva del documento word
                        document.add_page_break() 

        #elimina i file nella cartella per poi eliminare la cartella stessa
        folder = os.path.expanduser("~/Desktop/Immagini Temporanee/")
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print(f'Failed to delete {file_path}. Reason: {e}')
        os.rmdir(folder)
        document.save(os.path.expanduser(f'~/Desktop/Prodotti_{ricerca.replace(" ", "_")}.docx'))

        
    ###########################################################################################################################



    elif scelta_def == 2:

        HEADERS = ({'User-Agent':
                'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:106.0) Gecko/20100101 Firefox/106.0'})

        print("Vuoi vedere la lista di prodotti di cui stai tenendo traccia o vuoi inserire un nuovo prodotto?")

        scelta_2 = int(input("Per inserire un prodotto premi 1, per visualizzare la lista dei prodotti premi 2, per eliminare un prodotto premi 3: "))



        #inserisco un nuovo prodotto tramite link
        if scelta_2 == 1:
            
            prodotto = input("Inserire il link del prodotto di cui si vuole tenere traccia: ")
            webpage = requests.get(prodotto, headers=HEADERS)
            soup = BeautifulSoup(webpage.text, "html.parser")

            block_div_prezzo = soup.find_all("span", attrs={"class":"a-price aok-align-center reinventPricePriceToPayMargin priceToPay"})
            block_div_title = soup.find_all("div", attrs={"class":"a-section a-spacing-none"})
            block_div_prezzo2 = soup.find_all("span", attrs={"class":"a-size-base a-color-price a-color-price"})
            
            
            #cerca il prezzo
            if "a-price-whole" in str(block_div_prezzo):
                for block in range(len(block_div_prezzo)):
                    prezzo_temp = str(block_div_prezzo[block].find("span", attrs={'class':'a-offscreen'})).replace("<span class=\"a-offscreen\">", "").replace("</span>", "").replace("€", "").replace(".", "")
                    prezzo_fin = float(prezzo_temp.replace(",", "."))
            elif "a-size-base a-color-price a-color-price" in str(block_div_prezzo2):
                for block in range(len(block_div_prezzo2)):            
                    prezzo_temp = str(block_div_prezzo2[block]).replace("<span class=\"a-size-base a-color-price a-color-price\">", "").replace("</span>", "").replace("€", "").replace(".", "").replace("&nbsp", "")       
                    prezzo_fin = float(prezzo_temp.replace(",", "."))
            elif "a-spacing-none a-text-left a-size-mini twisterSwatchPrice" in  str(block_div_prezzo2):
                prezzo_temp = str(block_div_prezzo2[block]).replace("<span class=\"a-spacing-none a-text-left a-size-mini twisterSwatchPrice\">", "").replace("</span>", "").replace("€", "").replace(".", "").replace("&nbsp", "")       
                prezzo_fin = float(prezzo_temp.replace(",", "."))
            else:
                print("prezzo non trovato")

            #cerca il nome del prodotto
            for block in range(len(block_div_title)):        
                title = soup.find("span", attrs={"id":"productTitle"}).text.strip() 
            
            filename = 'products.csv'
            with open(filename, 'a') as csv_file:
                pass # Nessuna operazione effettuata

            with open(filename, 'r') as csv_file:
                reader = csv.reader(csv_file)
                row_num=0
                for row in reader:
                    row_num +=1           


            lista_prodotti = [(str(row_num+1)+")",title, str(prezzo_fin), prodotto)]
            
            with open("products.csv", "a", newline="", encoding="utf-8") as csvfile:
                writer = csv.writer(csvfile)
                
                for product in lista_prodotti:
                    writer.writerow(product)
                    


        #leggo il file csv con i prodotti e controlla se i prezzi sono diminuiti
        elif scelta_2 == 2:
            
                       
            
            with open("products.csv", "r") as csvfile:
                reader = csv.reader(csvfile)
                for row in reader:
                    row_num, product_name, price, link = row
                    print(f"{row_num} Nome prodotto: {product_name}, Prezzo: {price}, Link: {link} \n")
                    prodotto = link
                    webpage = requests.get(prodotto, headers=HEADERS)
                    soup = BeautifulSoup(webpage.text, "html.parser")
                    block_div_prezzo = soup.find_all("span", attrs={"class":"a-price aok-align-center reinventPricePriceToPayMargin priceToPay"})
                    block_div_prezzo2 = soup.find_all("span", attrs={"class":"a-size-base a-color-price a-color-price"})

                    if "a-price-whole" in str(block_div_prezzo):
                        for block in range(len(block_div_prezzo)):
                            prezzo_temp2 = str(block_div_prezzo[block].find("span", attrs={'class':'a-offscreen'})).replace("<span class=\"a-offscreen\">", "").replace("</span>", "").replace("€", "").replace(".", "")
                            prezzo_fin2 = float(prezzo_temp2.replace(",", "."))
                    elif "a-size-base a-color-price a-color-price" in str(block_div_prezzo2):
                        for block in range(len(block_div_prezzo2)):            
                            prezzo_temp2 = str(block_div_prezzo2[block]).replace("<span class=\"a-size-base a-color-price a-color-price\">", "").replace("</span>", "").replace("€", "").replace(".", "").replace("&nbsp", "")       
                            prezzo_fin2 = float(prezzo_temp2.replace(",", "."))
                    
                    if prezzo_fin2 < float(price):
                        
                        winsound.PlaySound("SystemExclamation", winsound.SND_ALIAS)
                        print(f"Il prezzo del prodotto appena mostrato è minore del prezzo di quando lo hai inserito, ora vale: {prezzo_fin2}\n")

        elif scelta_2== 3:
            
            # Definisci il nome del file CSV
            filename = 'products.csv'

            # Definisci l'indice della riga che vuoi rimuovere
            row_index = int(input('Inserire la riga da eliminare: '))

            # Crea un file temporaneo
            temp_file = tempfile.NamedTemporaryFile(mode='w', newline='', delete=False)

            # Apri il file CSV e il file temporaneo
            with open(filename, 'r') as csv_file, temp_file:
                # Crea un oggetto lettore CSV e un oggetto scrittore CSV
                reader = csv.reader(csv_file)
                writer = csv.writer(temp_file)

                # Itera le righe del file di input CSV e scrivi le righe tranne quella che vuoi rimuovere nel file temporaneo
                for i, row in enumerate(reader,start=1):
                    if i != row_index:
                        writer.writerow(row)

            # Sostituisci il file CSV esistente con il file temporaneo
            shutil.move(temp_file.name, filename)



    elif scelta_def==0:
        break

    else:
        print("Scelta non valida")
        continue











    